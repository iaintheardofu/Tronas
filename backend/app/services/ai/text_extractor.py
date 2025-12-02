"""
Text extraction service for various document types.
Supports PDF, Word, Excel, images (OCR), and email files.
"""
from typing import Optional, Dict, Any, List, BinaryIO
from pathlib import Path
import hashlib
import tempfile
import asyncio
from concurrent.futures import ThreadPoolExecutor

from loguru import logger

from app.core.config import settings


class TextExtractor:
    """
    Extract text content from various document types.
    Supports PDF, Word, Excel, PowerPoint, images, and email formats.
    """

    def __init__(self, max_workers: int = 4):
        self.executor = ThreadPoolExecutor(max_workers=max_workers)

    async def extract_text(
        self,
        file_path: str = None,
        file_content: bytes = None,
        mime_type: str = None,
        filename: str = None,
    ) -> Dict[str, Any]:
        """
        Extract text from a document.

        Args:
            file_path: Path to the file
            file_content: Raw file bytes
            mime_type: MIME type of the file
            filename: Original filename (for type detection)

        Returns:
            Dictionary with extracted text and metadata
        """
        # Read file content if path provided
        if file_path and not file_content:
            with open(file_path, "rb") as f:
                file_content = f.read()
            if not filename:
                filename = Path(file_path).name

        if not file_content:
            raise ValueError("Either file_path or file_content must be provided")

        # Detect mime type if not provided
        if not mime_type:
            mime_type = self._detect_mime_type(file_content, filename)

        # Calculate hash
        file_hash = hashlib.sha256(file_content).hexdigest()

        # Extract based on type
        loop = asyncio.get_event_loop()

        try:
            if mime_type == "application/pdf":
                text, metadata = await loop.run_in_executor(
                    self.executor, self._extract_pdf, file_content
                )
            elif mime_type in [
                "application/msword",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ]:
                text, metadata = await loop.run_in_executor(
                    self.executor, self._extract_word, file_content
                )
            elif mime_type in [
                "application/vnd.ms-excel",
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            ]:
                text, metadata = await loop.run_in_executor(
                    self.executor, self._extract_excel, file_content
                )
            elif mime_type in ["message/rfc822"]:
                text, metadata = await loop.run_in_executor(
                    self.executor, self._extract_eml, file_content
                )
            elif mime_type == "application/vnd.ms-outlook":
                text, metadata = await loop.run_in_executor(
                    self.executor, self._extract_msg, file_content
                )
            elif mime_type.startswith("image/"):
                text, metadata = await loop.run_in_executor(
                    self.executor, self._extract_image_ocr, file_content
                )
            elif mime_type == "text/plain":
                text = file_content.decode("utf-8", errors="replace")
                metadata = {"type": "text"}
            elif mime_type == "text/html":
                text, metadata = await loop.run_in_executor(
                    self.executor, self._extract_html, file_content
                )
            else:
                # Try as plain text
                try:
                    text = file_content.decode("utf-8", errors="replace")
                    metadata = {"type": "unknown_text"}
                except Exception:
                    text = ""
                    metadata = {"type": "unsupported", "error": "Could not extract text"}

            return {
                "text": text,
                "file_hash": file_hash,
                "mime_type": mime_type,
                "filename": filename,
                "metadata": metadata,
                "char_count": len(text),
                "word_count": len(text.split()) if text else 0,
            }

        except Exception as e:
            logger.error(f"Text extraction failed for {filename}: {e}")
            return {
                "text": "",
                "file_hash": file_hash,
                "mime_type": mime_type,
                "filename": filename,
                "metadata": {"error": str(e)},
                "char_count": 0,
                "word_count": 0,
            }

    def _detect_mime_type(self, content: bytes, filename: str = None) -> str:
        """Detect MIME type from content and filename."""
        try:
            import magic
            mime = magic.from_buffer(content, mime=True)
            return mime
        except ImportError:
            # Fallback to extension-based detection
            if filename:
                ext = Path(filename).suffix.lower()
                mime_map = {
                    ".pdf": "application/pdf",
                    ".doc": "application/msword",
                    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    ".xls": "application/vnd.ms-excel",
                    ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    ".msg": "application/vnd.ms-outlook",
                    ".eml": "message/rfc822",
                    ".txt": "text/plain",
                    ".html": "text/html",
                    ".jpg": "image/jpeg",
                    ".jpeg": "image/jpeg",
                    ".png": "image/png",
                }
                return mime_map.get(ext, "application/octet-stream")
            return "application/octet-stream"

    def _extract_pdf(self, content: bytes) -> tuple[str, Dict[str, Any]]:
        """Extract text from PDF using pypdf."""
        from pypdf import PdfReader
        from io import BytesIO

        reader = PdfReader(BytesIO(content))
        text_parts = []
        metadata = {
            "type": "pdf",
            "page_count": len(reader.pages),
        }

        # Extract document info
        if reader.metadata:
            metadata["title"] = reader.metadata.get("/Title", "")
            metadata["author"] = reader.metadata.get("/Author", "")
            metadata["subject"] = reader.metadata.get("/Subject", "")
            metadata["creator"] = reader.metadata.get("/Creator", "")

        # Extract text from each page
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(f"[Page {i + 1}]\n{page_text}")

        return "\n\n".join(text_parts), metadata

    def _extract_word(self, content: bytes) -> tuple[str, Dict[str, Any]]:
        """Extract text from Word documents."""
        from docx import Document
        from io import BytesIO

        doc = Document(BytesIO(content))
        text_parts = []
        metadata = {"type": "word"}

        # Extract core properties
        if doc.core_properties:
            metadata["title"] = doc.core_properties.title or ""
            metadata["author"] = doc.core_properties.author or ""
            metadata["created"] = str(doc.core_properties.created) if doc.core_properties.created else ""
            metadata["modified"] = str(doc.core_properties.modified) if doc.core_properties.modified else ""

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Extract tables
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                table_text.append(" | ".join(row_text))
            if table_text:
                text_parts.append("[TABLE]\n" + "\n".join(table_text))

        return "\n\n".join(text_parts), metadata

    def _extract_excel(self, content: bytes) -> tuple[str, Dict[str, Any]]:
        """Extract text from Excel files."""
        from openpyxl import load_workbook
        from io import BytesIO

        wb = load_workbook(BytesIO(content), data_only=True)
        text_parts = []
        metadata = {
            "type": "excel",
            "sheet_count": len(wb.sheetnames),
            "sheets": wb.sheetnames,
        }

        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            sheet_text = [f"[Sheet: {sheet_name}]"]

            for row in sheet.iter_rows(values_only=True):
                row_values = [str(cell) if cell is not None else "" for cell in row]
                if any(row_values):
                    sheet_text.append(" | ".join(row_values))

            text_parts.append("\n".join(sheet_text))

        return "\n\n".join(text_parts), metadata

    def _extract_eml(self, content: bytes) -> tuple[str, Dict[str, Any]]:
        """Extract text from EML email files."""
        import email
        from email import policy

        msg = email.message_from_bytes(content, policy=policy.default)

        metadata = {
            "type": "eml",
            "subject": msg.get("Subject", ""),
            "from": msg.get("From", ""),
            "to": msg.get("To", ""),
            "cc": msg.get("Cc", ""),
            "date": msg.get("Date", ""),
        }

        text_parts = [
            f"From: {metadata['from']}",
            f"To: {metadata['to']}",
            f"Subject: {metadata['subject']}",
            f"Date: {metadata['date']}",
            "",
        ]

        # Extract body
        if msg.is_multipart():
            for part in msg.walk():
                content_type = part.get_content_type()
                if content_type == "text/plain":
                    text_parts.append(part.get_payload(decode=True).decode("utf-8", errors="replace"))
        else:
            text_parts.append(msg.get_payload(decode=True).decode("utf-8", errors="replace"))

        return "\n".join(text_parts), metadata

    def _extract_msg(self, content: bytes) -> tuple[str, Dict[str, Any]]:
        """Extract text from Outlook MSG files."""
        import extract_msg
        from io import BytesIO
        import tempfile
        import os

        # extract_msg requires a file, so write temporarily
        with tempfile.NamedTemporaryFile(delete=False, suffix=".msg") as tmp:
            tmp.write(content)
            tmp_path = tmp.name

        try:
            msg = extract_msg.Message(tmp_path)

            metadata = {
                "type": "msg",
                "subject": msg.subject or "",
                "from": msg.sender or "",
                "to": msg.to or "",
                "cc": msg.cc or "",
                "date": str(msg.date) if msg.date else "",
            }

            text_parts = [
                f"From: {metadata['from']}",
                f"To: {metadata['to']}",
                f"Subject: {metadata['subject']}",
                f"Date: {metadata['date']}",
                "",
                msg.body or "",
            ]

            return "\n".join(text_parts), metadata

        finally:
            os.unlink(tmp_path)

    def _extract_html(self, content: bytes) -> tuple[str, Dict[str, Any]]:
        """Extract text from HTML files."""
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(content, "html.parser")

        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()

        text = soup.get_text(separator="\n")
        # Clean up whitespace
        lines = (line.strip() for line in text.splitlines())
        text = "\n".join(line for line in lines if line)

        metadata = {
            "type": "html",
            "title": soup.title.string if soup.title else "",
        }

        return text, metadata

    def _extract_image_ocr(self, content: bytes) -> tuple[str, Dict[str, Any]]:
        """Extract text from images using OCR (placeholder - requires additional setup)."""
        # Note: Full OCR would require pytesseract or Azure Computer Vision
        # This is a placeholder that would need cloud OCR integration

        metadata = {
            "type": "image",
            "ocr_note": "OCR processing requires additional configuration"
        }

        # Return empty text - would integrate with Azure Computer Vision for production
        return "", metadata

    async def extract_batch(
        self,
        files: List[Dict[str, Any]],
        concurrency: int = 4,
    ) -> List[Dict[str, Any]]:
        """
        Extract text from multiple files concurrently.

        Args:
            files: List of file dicts with 'content' or 'path', and optional 'filename', 'mime_type'
            concurrency: Number of concurrent extractions

        Returns:
            List of extraction results
        """
        semaphore = asyncio.Semaphore(concurrency)
        results = []

        async def extract_with_limit(file_info: Dict[str, Any], index: int):
            async with semaphore:
                result = await self.extract_text(
                    file_path=file_info.get("path"),
                    file_content=file_info.get("content"),
                    mime_type=file_info.get("mime_type"),
                    filename=file_info.get("filename"),
                )
                result["index"] = index
                result["id"] = file_info.get("id")
                return result

        tasks = [extract_with_limit(f, i) for i, f in enumerate(files)]
        results = await asyncio.gather(*tasks, return_exceptions=True)

        # Handle exceptions
        processed = []
        for i, result in enumerate(results):
            if isinstance(result, Exception):
                processed.append({
                    "text": "",
                    "error": str(result),
                    "index": i,
                })
            else:
                processed.append(result)

        return processed


# Singleton instance
_extractor: Optional[TextExtractor] = None


def get_text_extractor() -> TextExtractor:
    """Get or create the text extractor singleton."""
    global _extractor
    if _extractor is None:
        _extractor = TextExtractor()
    return _extractor
